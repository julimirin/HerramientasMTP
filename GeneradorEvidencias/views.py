#!/usr/bin/env python
# -*- coding: utf-8 -*-
from django.shortcuts import render, redirect  # puedes importar render_to_response
from .forms import GenerarForm
from .models import  Cliente, Solicitud, CasoPrueba, PasoPrueba, Plantilla
from openpyxl.utils import get_column_letter
import openpyxl
from docxtpl import DocxTemplate
from docx import Document
import zipfile
import os, shutil
import time
from django.conf import settings
import unicodedata

my_media_root = settings.MEDIA_ROOT


def elimina_tildes(s):
   return ''.join((c for c in unicodedata.normalize('NFD', s) if unicodedata.category(c) != 'Mn'))

def generar_evidencias(request):
    if request.method == 'POST':
        form = GenerarForm(request.POST, request.FILES)
        if form.is_valid():
            nueva_solicitud = form.save()
            solicitud = Solicitud.objects.get(id=nueva_solicitud.pk)
            insertar_solicitud(solicitud)
            generar_documentos(solicitud)
            file = generar_documentos_zip(solicitud)
            #return redirect('/GeneradorEvidencias/generar_descarga')
            return render(request, 'GeneradorEvidencias/generador_descarga.html', {'file': file})
    else:
        form = GenerarForm()
    return render(request, 'GeneradorEvidencias/generador_evidencias.html', {'form': form})

def insertar_solicitud(solicitud):
    # Abrimos el fichero indicado en el archivo de DatosProyecto.xlsx

    doc = openpyxl.load_workbook(solicitud.archivo)


    #Abrimos la hoja correspondiente en dicho fichero , en Repsol es DEFPRU (FUN)
    hoja = doc.get_sheet_by_name('CASOS')

    # Guardamos las claves de los datos del plan de pruebas (cabeceras)
    sel_claves = hoja['A1':'F1']
    claves = []
    for fila in sel_claves:
        for columna in fila:
            claves.append(columna.value)

    # Guardamos los datos asociados a las claves por cada caso de prueba (valores)
    inicio = 'A2'
    fin = 'F' + str(hoja.max_row)
    seleccion = hoja[inicio:fin]
    testid = ""
    for fila in seleccion:
        for columna in fila:
            if columna.coordinate.startswith("A"):
                oldtestid = testid
                testid = columna.value
                # Es el mismo caso, añadimos los step que vengan al caso
                if testid == oldtestid:
                    paso = {}
                    for columna in fila:
                        if columna.coordinate.startswith("A"):
                            paso['codigo_caso'] = columna.value
                        elif columna.coordinate.startswith("D"):
                            paso['numero_paso'] = columna.value
                        elif columna.coordinate.startswith("E"):
                            paso['descripcion_paso'] = columna.value
                        elif columna.coordinate.startswith("F"):
                            paso['resultado_paso'] = columna.value

                    cprueba_relacionado = CasoPrueba.objects.get(codigo_caso = paso['codigo_caso'], solicitud=solicitud.pk)

                    nuevo_paso = PasoPrueba(
                        codigo_caso = cprueba_relacionado,
                        numero_paso = paso['numero_paso'],
                        descripcion_paso = paso['descripcion_paso'],
                        resultado_paso = paso['resultado_paso']
                    )
                    nuevo_paso.save()
                # Se trata de un caso nuevo, lo creamos
                else:
                    caso = {}
                    paso = {}
                    for columna in fila:
                        if columna.coordinate.startswith("A"):
                            caso['codigo_caso'] = columna.value
                        elif columna.coordinate.startswith("B"):
                            caso['nombre_caso'] = columna.value
                        elif columna.coordinate.startswith("C"):
                            caso['descripcion_caso'] = columna.value
                        elif columna.coordinate.startswith("D"):
                            paso['numero_paso'] = columna.value
                        elif columna.coordinate.startswith("E"):
                            paso['descripcion_paso'] = columna.value
                        elif columna.coordinate.startswith("F"):
                            paso['resultado_paso'] = columna.value

                    nuevo_caso = CasoPrueba(
                        solicitud = Solicitud.objects.get(id=solicitud.pk),
                        codigo_caso = caso['codigo_caso'],
                        nombre_caso = caso['nombre_caso'],
                        descripcion_caso =  caso['descripcion_caso']
                    )
                    nuevo_caso.save()

                    cprueba_relacionado = CasoPrueba.objects.get(codigo_caso=caso['codigo_caso'], solicitud=solicitud.pk)

                    nuevo_paso = PasoPrueba(
                        codigo_caso=cprueba_relacionado,
                        numero_paso=paso['numero_paso'],
                        descripcion_paso=paso['descripcion_paso'],
                        resultado_paso=paso['resultado_paso']
                    )
                    nuevo_paso.save()

def generar_documentos(solicitud):



    for caso_prueba in CasoPrueba.objects.filter(solicitud = solicitud.pk) :
        plantilla = Plantilla.objects.get(cliente=solicitud.cliente)
        tpl = DocxTemplate(plantilla.plantilla)
        print(caso_prueba.codigo_caso, caso_prueba.nombre_caso)
        context = {
            'sprint': '',
            'proyecto': solicitud.codigo_proyecto,
            'nombre_proyecto': solicitud.nombre_proyecto,
            'entorno': solicitud.entorno,
            'fase_pruebas': solicitud.fase_de_prueba,
            'codigo_caso': caso_prueba.codigo_caso,
            'nombre_caso': caso_prueba.nombre_caso
        }
        tpl.render(context)
        tpl.save(my_media_root+"/evidencias/"+ caso_prueba.codigo_caso + " " + caso_prueba.nombre_caso + '.docx')
        for paso in PasoPrueba.objects.filter(codigo_caso = caso_prueba):
            print(caso_prueba.codigo_caso,paso.numero_paso)
            document = Document(my_media_root +"/evidencias/"+ caso_prueba.codigo_caso + " " + caso_prueba.nombre_caso + '.docx')
            p = document.add_paragraph()
            p.add_run(paso.numero_paso).bold = True
            p.add_run(': ').bold = True
            p.add_run(paso.descripcion_paso)
            p = document.add_paragraph('')
            p.add_run('Resultado Esperado: ').italic = True
            p.add_run(paso.resultado_paso).italic = True
            document.save(my_media_root+"/evidencias/"+ caso_prueba.codigo_caso + " " + elimina_tildes(caso_prueba.nombre_caso) + '.docx')

def generar_documentos_zip(solicitud):

    nombre_archivo = solicitud.codigo_proyecto + time.strftime("%Y%m%d-%H%M%S") + '.zip'
    documentos_zip = zipfile.ZipFile(my_media_root+'/evidencias/'+ nombre_archivo , 'w')

    for folder, subfolders, files in os.walk(my_media_root+'/evidencias/'):
        for file in files:
            if file.endswith('.docx'):
                documentos_zip.write(os.path.join(folder, file),
                                  os.path.relpath(os.path.join(folder, file), my_media_root+'/evidencias/'),
                                  compress_type=zipfile.ZIP_DEFLATED)

    documentos_zip.close()

    #Borramos los word
    directory = my_media_root+'/evidencias/'
    test = os.listdir(directory)

    for item in test:
        if item.endswith(".docx"):
            os.remove(os.path.join(directory, item))


    return nombre_archivo


def generar_descarga(request):
    return render(request, 'GeneradorEvidencias/generador_descarga.html')









